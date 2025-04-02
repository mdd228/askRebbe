import os
import unittest
from models.document import Document
from controllers.document_controller import DocumentController
from docx import Document as DocxDocument
import win32com.client
import pythoncom
import shutil
import time

class TestDocumentProcessing(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        """Set up test environment once for all tests"""
        cls.base_test_dir = os.path.join(os.path.dirname(__file__), 'test_files')
        if os.path.exists(cls.base_test_dir):
            shutil.rmtree(cls.base_test_dir)
        os.makedirs(cls.base_test_dir)
    
    def setUp(self):
        """Set up test environment for each test"""
        # Create a unique test directory for this test
        self.test_dir = os.path.join(self.base_test_dir, f'test_{time.time_ns()}')
        os.makedirs(self.test_dir)
        
        # Create test content
        self.test_content = "This is a test document.\nIt contains multiple lines.\nFor testing purposes."
        
        # Create test files
        self.test_files = {}
        
        # Create TXT file
        self.test_files['txt'] = os.path.join(self.test_dir, 'test.txt')
        with open(self.test_files['txt'], 'w', encoding='utf-8') as f:
            f.write(self.test_content)
        self._wait_for_file(self.test_files['txt'])
        
        # Create DOCX file
        self.test_files['docx'] = os.path.join(self.test_dir, 'test.docx')
        self._create_docx_file(self.test_files['docx'])
        
        # Create DOC file
        self.test_files['doc'] = os.path.join(self.test_dir, 'test.doc')
        self._create_doc_file(self.test_files['doc'])
            
        # Initialize controller with absolute path
        self.controller = DocumentController(os.path.abspath(self.test_dir))
        
        # Verify all files exist before proceeding
        for file_path in self.test_files.values():
            self.assertTrue(os.path.exists(file_path), f"File not created: {file_path}")
    
    def _create_docx_file(self, file_path):
        """Create a test DOCX file"""
        doc = DocxDocument()
        doc.add_paragraph(self.test_content)
        doc.save(file_path)
        # Ensure file is written
        self._wait_for_file(file_path)
    
    def _create_doc_file(self, file_path):
        """Create a test DOC file using Word COM"""
        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Add()
            doc.Content.Text = self.test_content
            doc.SaveAs(file_path)
            doc.Close()
            word.Quit()
            pythoncom.CoUninitialize()
            # Ensure file is written
            self._wait_for_file(file_path)
        except Exception as e:
            print(f"Error creating DOC file: {str(e)}")
    
    def _wait_for_file(self, file_path, timeout=5):
        """Wait for file to be fully written and accessible"""
        start_time = time.time()
        while not os.path.exists(file_path) and time.time() - start_time < timeout:
            time.sleep(0.1)
        if os.path.exists(file_path):
            # Additional wait to ensure file is fully written
            time.sleep(0.5)
    
    def test_document_creation(self):
        """Test basic document creation and file type detection"""
        for ext, file_path in self.test_files.items():
            self.assertTrue(os.path.exists(file_path), f"File not found: {file_path}")
            doc = Document(file_path)
            self.assertEqual(doc.filename, f'test.{ext}')
            self.assertEqual(doc.file_type, ext)
    
    def test_directory_scanning(self):
        """Test if controller can find files in directory"""
        files = self.controller.scan_directory()
        for file_path in self.test_files.values():
            self.assertIn(os.path.abspath(file_path), files)
    
    def test_content_extraction(self):
        """Test if content can be extracted correctly from all file types"""
        for ext, file_path in self.test_files.items():
            self.assertTrue(os.path.exists(file_path), f"File not found: {file_path}")
            
            # Test through Document model
            doc = Document(file_path)
            content = doc.extract_content()
            self.assertIsNotNone(content, f"Failed to extract content from {ext} file")
            self.assertIn(self.test_content, content, f"Content mismatch in {ext} file")
            
            # Test through controller
            controller_content = self.controller.process_document(os.path.abspath(file_path))
            
            # If this is a DOCX file and a corresponding DOC exists, it should be skipped
            if ext == 'docx' and os.path.exists(file_path.replace('.docx', '.doc')):
                self.assertIsNone(controller_content, 
                                f"DOCX file should be skipped when DOC exists: {file_path}")
            else:
                self.assertIsNotNone(controller_content, 
                                   f"Controller failed to process {ext} file")
                self.assertIn(self.test_content, controller_content, 
                            f"Controller content mismatch in {ext} file")
    
    def test_controller_batch_processing(self):
        """Test if controller can process all documents at once"""
        # Verify all files exist before processing
        for file_path in self.test_files.values():
            self.assertTrue(os.path.exists(file_path), f"File not found before batch processing: {file_path}")
        
        results = self.controller.process_all_documents()
        # We expect results for all original files except the DOCX file when a DOC exists
        expected_count = len(self.test_files)
        if 'doc' in self.test_files and 'docx' in self.test_files:
            expected_count -= 1  # Subtract 1 because DOCX will be skipped when DOC exists
            
        self.assertEqual(len(results), expected_count, 
                        f"Expected {expected_count} results, got {len(results)}")
        
        for file_path, content in results.items():
            self.assertIsNotNone(content, f"Failed to process {file_path}")
            self.assertIn(self.test_content, content, f"Content mismatch in {file_path}")
            
            # If this is a DOCX file, ensure it's not a temporary one
            if file_path.endswith('.docx'):
                doc_path = file_path.replace('.docx', '.doc')
                self.assertFalse(os.path.exists(doc_path), 
                               f"DOCX file {file_path} should not be processed when DOC exists")
    
    def tearDown(self):
        """Clean up test files for this test"""
        # Add a small delay before cleanup to ensure all file operations are complete
        time.sleep(0.5)
        if os.path.exists(self.test_dir):
            shutil.rmtree(self.test_dir)
    
    @classmethod
    def tearDownClass(cls):
        """Clean up the entire test environment"""
        if os.path.exists(cls.base_test_dir):
            shutil.rmtree(cls.base_test_dir)

if __name__ == '__main__':
    unittest.main() 