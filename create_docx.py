from docx import Document

# Create a new document
doc = Document()

# Add content
doc.add_paragraph("The Lubavitcher Rebbe's Teachings on Jewish Education")
doc.add_paragraph("The Rebbe emphasized the importance of Jewish education as a fundamental pillar of Jewish life. His approach to education was unique and comprehensive, focusing on both spiritual and practical aspects.")
doc.add_paragraph("Key principles of The Rebbe's educational philosophy:")
doc.add_paragraph("1. Individual Connection: Each student must be treated as a unique individual with their own strengths and challenges.")
doc.add_paragraph("2. Love for Learning: Education should instill a love for Torah and Jewish learning.")
doc.add_paragraph("3. Practical Application: Learning must be connected to real-life situations and practical mitzvah observance.")
doc.add_paragraph("4. Spiritual Growth: Education should focus on both intellectual development and spiritual growth.")
doc.add_paragraph("5. Community Involvement: Learning should take place within a supportive Jewish community.")

# Save the document
doc.save('pdfs/test_docx.docx') 