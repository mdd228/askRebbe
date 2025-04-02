from typing import Optional
from docx import Document as DocxDocument
import PyPDF2
import os
import pythoncom
import win32com.client
from contextlib import contextmanager

class Document:
    def __init__(self, file_path: str):
        self.file_path = os.path.abspath(file_path)
        self.filename = os.path.basename(file_path)
        self.content: Optional[str] = None
        self.file_type = self._get_file_type()
        
    def _get_file_type(self) -> str:
        """Determine the file type based on extension."""
        ext = self.filename.lower().split('.')[-1]
        return ext
    
    @contextmanager
    def _word_app(self):
        """Context manager for Word application to ensure proper cleanup."""
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            yield word
        finally:
            try:
                word.Quit()
            except:
                pass
            pythoncom.CoUninitialize()
    
    def extract_content(self) -> Optional[str]:
        """Extract content from the document based on its type."""
        if not os.path.exists(self.file_path):
            print(f"File does not exist: {self.file_path}")
            return None
            
        try:
            if self.file_type == 'pdf':
                self.content = self._extract_pdf()
            elif self.file_type == 'docx':
                self.content = self._extract_docx()
            elif self.file_type == 'doc':
                self.content = self._extract_doc()
            elif self.file_type == 'txt':
                self.content = self._extract_txt()
            return self.content
        except Exception as e:
            print(f"Error extracting content from {self.filename}: {str(e)}")
            return None
    
    def _extract_txt(self) -> Optional[str]:
        """Extract text from a text file."""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error processing text file {self.file_path}: {str(e)}")
            return None
    
    def _extract_pdf(self) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            with open(self.file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                return '\n'.join(text)
        except Exception as e:
            print(f"Error processing PDF file {self.file_path}: {str(e)}")
            return None
    
    def _extract_docx(self) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(self.file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text.append(paragraph.text)
            result = '\n'.join(text)
            return result
        except Exception as e:
            print(f"Error processing DOCX file {self.file_path}: {str(e)}")
            return None
    
    def _extract_doc(self) -> Optional[str]:
        """Extract text from DOC file by converting to DOCX first."""
        try:
            # Convert to DOCX
            docx_path = self._convert_doc_to_docx()
            if docx_path and os.path.exists(docx_path):
                try:
                    # Create temporary Document instance for DOCX
                    temp_doc = Document(docx_path)
                    content = temp_doc.extract_content()
                    return content
                finally:
                    # Clean up the temporary DOCX file
                    if os.path.exists(docx_path):
                        os.remove(docx_path)
            return None
        except Exception as e:
            print(f"Error processing DOC file {self.file_path}: {str(e)}")
            return None
    
    def _convert_doc_to_docx(self) -> Optional[str]:
        """Convert DOC to DOCX using Word COM automation."""
        docx_path = self.file_path.replace('.doc', '.docx')
        
        try:
            with self._word_app() as word:
                doc = word.Documents.Open(self.file_path)
                doc.SaveAs2(docx_path, FileFormat=16)  # 16 = docx format
                doc.Close()
            return docx_path
        except Exception as e:
            print(f"Error converting DOC to DOCX {self.file_path}: {str(e)}")
            if os.path.exists(docx_path):
                os.remove(docx_path)
            return None

    def to_dict(self) -> dict:
        """Convert document to dictionary representation."""
        return {
            'filename': self.filename,
            'content': self.content if self.content else ''
        } 